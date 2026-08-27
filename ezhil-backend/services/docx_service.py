"""
DOCX Text Extraction Service
=============================
Extracts readable text from .docx Word documents, preserving:
  - Paragraph text (body, headings)
  - Table cell content (row-by-row, tab-separated)
  - Tamil Unicode characters are preserved as-is

Dependencies:
    pip install python-docx
"""
from __future__ import annotations

import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text(docx_bytes: bytes) -> Tuple[str, int]:
    """
    Extract plain text from a .docx file.

    Returns:
        (text, paragraph_count)
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed — pip install python-docx")
        return "", 0

    try:
        doc      = Document(io.BytesIO(docx_bytes))
        parts: list[str] = []

        # Body paragraphs
        para_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
                para_count += 1

        # Tables — flatten each row as tab-separated cells
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        combined = "\n".join(parts)
        logger.info("DOCX extraction: %d paragraphs, %d chars", para_count, len(combined))
        return combined, para_count

    except Exception as exc:
        logger.error("DOCX extraction error: %s", exc)
        return "", 0
